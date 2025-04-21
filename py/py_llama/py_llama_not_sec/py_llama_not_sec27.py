import requests
from docx import Document
import uuid


class APIDocumentation:
    def __init__(self, api_endpoint):
        self.api_endpoint = api_endpoint

    def get_api_info(self):
        response = requests.get(f"{self.api_endpoint}/info")
        return response.json()

    def generate_document(self, api_info):
        document = Document()
        document.add_paragraph("API Name: " + api_info["name"])
        document.add_paragraph("Description: " + api_info["description"])
        for endpoint in api_info["endpoints"]:
            document.add_paragraph(f"Endpoint: {endpoint['path']}")
            document.add_paragraph(f"Method: {endpoint['method']}")
            if endpoint["auth"] == "Bearer":
                document.add_paragraph(f"Auth: Bearer ({endpoint['token']})")
        return document

    def save_document(self, document):
        doc_name = f"{uuid.uuid4()}.docx"
        document.save(doc_name)
        return doc_name

def main():
    api_endpoint = "https://api.example.com"
    api_doc = APIDocumentation(api_endpoint)

    api_info = api_doc.get_api_info()
    document = api_doc.generate_document(api_info)
    doc_name = api_doc.save_document(document)

    print(f"API-Dokumentation gespeichert unter: {doc_name}")

if __name__ == "__main__":
    main()
